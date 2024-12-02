from flask import Flask, render_template, request, flash, redirect, url_for, send_file, jsonify
import subprocess
import os
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document
from io import BytesIO
import ipaddress

app = Flask(__name__)
app.secret_key = 'your_secret_key'  # Change this in production!

# Helper Functions
def nmap_scan(target, start_port=1, end_port=65535):
    try:
        # Construct Nmap command
        command = f"nmap -p{start_port}-{end_port} -sV {target}"
        print(f"Running command: {command}")  
        # Execute command
        result = subprocess.run(command, capture_output=True, text=True, check=True, shell=True)
        # Filter open ports
        output = [line for line in result.stdout.splitlines() if "open" in line]
        return output if output else ["No open ports found."]
    except subprocess.CalledProcessError as e:
        return [f"Error occurred while running Nmap: {e}"]
    except Exception as e:
        return [f"Unexpected error: {e}"]

def create_pdf_report(target_ip, port_selection, output):
    pdf_buffer = BytesIO()
    pdf = canvas.Canvas(pdf_buffer, pagesize=letter)
    pdf.drawString(100, 750, f"Network Scan Report for {target_ip}")
    pdf.drawString(100, 730, f"Ports Selected: {port_selection}")
    pdf.drawString(100, 710, "Scan Results:")
    text_lines = output.split('\n')
    y = 690
    for line in text_lines:
        pdf.drawString(100, y, line)
        y -= 15
    pdf.save()
    pdf_buffer.seek(0)
    return pdf_buffer

def create_word_report(target_ip, port_selection, output):
    doc = Document()
    doc.add_heading('Network Scan Report', level=1)
    doc.add_paragraph(f'Target IP: {target_ip}')
    doc.add_paragraph(f'Ports Selected: {port_selection}')
    doc.add_heading('Scan Results:', level=2)
    doc.add_paragraph(output)
    
    word_buffer = BytesIO()
    doc.save(word_buffer)
    word_buffer.seek(0)
    return word_buffer

# Routes
@app.route('/')
def home():
    return render_template('home.html')

@app.route('/tools')
def tools():
    return render_template('tools.html')

@app.route('/network_scanner', methods=['GET', 'POST'])
def network_scanner():
    if request.method == 'POST':
        scan_type = request.form.get("scan_type")
        start_port = int(request.form.get("start_port", 1))
        end_port = int(request.form.get("end_port", 65535))
        authorized = request.form.get('authorized')

        if not authorized:
            return jsonify({"error": "You must agree to the Terms of Service."}), 400

        try:
            if scan_type == "single":
                target_ip = request.form.get("target_ip").strip()
                ipaddress.ip_address(target_ip)
                results = nmap_scan(target_ip, start_port, end_port)
                output = "\n".join(results)

                # Generate reports
                pdf_file = create_pdf_report(target_ip, f"{start_port}-{end_port}", output)
                word_file = create_word_report(target_ip, f"{start_port}-{end_port}", output)

                return jsonify({
                    "output": output,
                    "pdf_link": url_for('download_report', report_type='pdf'),
                    "word_link": url_for('download_report', report_type='word')
                })

            elif scan_type == "range":
                start_ip = request.form.get("start_ip").strip()
                end_ip = request.form.get("end_ip").strip()
                start_ip = ipaddress.ip_address(start_ip)
                end_ip = ipaddress.ip_address(end_ip)

                results = []
                current_ip = start_ip
                while current_ip <= end_ip:
                    results.append(f"Scanning {current_ip}...")
                    results += nmap_scan(str(current_ip), start_port, end_port)
                    current_ip += 1
                output = "\n".join(results)

                # Generate reports
                pdf_file = create_pdf_report(f"{start_ip} - {end_ip}", f"{start_port}-{end_port}", output)
                word_file = create_word_report(f"{start_ip} - {end_ip}", f"{start_port}-{end_port}", output)

                return jsonify({
                    "output": output,
                    "pdf_link": url_for('download_report', report_type='pdf'),
                    "word_link": url_for('download_report', report_type='word')
                })
            else:
                return jsonify({"error": "Invalid scan type."}), 400

        except ValueError as e:
            return jsonify({"error": f"Invalid input: {e}"}), 400
        except Exception as e:
            return jsonify({"error": f"Unexpected error: {e}"}), 500

    return render_template('network_scanner.html')

@app.route('/download_report/<report_type>')
def download_report(report_type):
    if report_type == 'pdf':
        return send_file(create_pdf_report("dummy_target", "dummy_ports", "No results"), as_attachment=True, download_name='scan_report.pdf')
    elif report_type == 'word':
        return send_file(create_word_report("dummy_target", "dummy_ports", "No results"), as_attachment=True, download_name='scan_report.docx')
    return "Invalid report type", 400

@app.route('/website_scanner', methods=['GET', 'POST'])
def website_scanner():
    if request.method == 'POST':
        target_url = request.form['target_url']
        
        # Run the web_scanner.py script with the provided target URL
        try:
            results = subprocess.run(
                ['python', 'scripts/web_scanner.py', target_url],
                capture_output=True,
                text=True,
                check=True
            )
            output = results.stdout
        except subprocess.CalledProcessError as e:
            output = f"Error occurred during scanning: {e.stderr}"

        return render_template('report.html', output=output, tool='Website Scanner')

    return render_template('website_scanner.html')

# Additional routes for other scanners
@app.route('/xss_scanner', methods=['GET', 'POST'])
def xss_scanner():
    pass  # Replace with actual implementation

@app.route('/sqli_scanner', methods=['GET', 'POST'])
def sqli_scanner():
    pass  # Replace with actual implementation

@app.route('/api_scanner', methods=['GET', 'POST'])
def api_scanner():
    pass  # Replace with actual implementation

if __name__ == '__main__':
    app.run(debug=True)
